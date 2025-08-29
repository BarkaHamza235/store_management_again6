from io import BytesIO, StringIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
import csv

def generate_sales_pdf(queryset):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Liste des ventes")
    y -= 30
    c.setFont("Helvetica", 10)
    cols = ["N° Facture", "Date", "Caissier", "Montant", "Statut"]
    x_positions = [50, 150, 260, 380, 460]
    for x, col in zip(x_positions, cols):
        c.drawString(x, y, col)
    y -= 20
    for sale in queryset:
        if y < 50:
            c.showPage()
            y = height - 50
        values = [
            sale.invoice_number,
            sale.date.strftime("%d/%m/%Y %H:%M"),
            sale.cashier.get_full_name(),
            f"{sale.total_amount:.2f}€",
            sale.get_status_display(),
        ]
        for x, val in zip(x_positions, values):
            c.drawString(x, y, val)
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer

def generate_sales_excel(queryset):
    wb = Workbook()
    ws = wb.active
    ws.title = "Ventes"
    ws.append(['N° Facture', 'Date', 'Caissier', 'Montant', 'Statut'])
    for sale in queryset:
        ws.append([
            sale.invoice_number,
            sale.date.strftime('%d/%m/%Y %H:%M'),
            sale.cashier.get_full_name(),
            sale.total_amount,
            sale.get_status_display(),
        ])
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

def generate_sales_docx(queryset):
    doc = Document()
    doc.add_heading('Liste des ventes', level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    for idx, title in enumerate(['N° Facture', 'Date', 'Caissier', 'Montant', 'Statut']):
        hdr_cells[idx].text = title
    for sale in queryset:
        row_cells = table.add_row().cells
        row_cells[0].text = sale.invoice_number
        row_cells[1].text = sale.date.strftime('%d/%m/%Y %H:%M')
        row_cells[2].text = sale.cashier.get_full_name()
        row_cells[3].text = f"{sale.total_amount:.2f}"
        row_cells[4].text = sale.get_status_display()
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream



def generate_sales_csv(queryset):
    # 1) Écrire dans un buffer texte
    text_stream = StringIO()
    writer = csv.writer(text_stream, delimiter='\t', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow(['N° Facture', 'Date', 'Caissier', 'Montant', 'Statut'])
    for sale in queryset:
        writer.writerow([
            sale.invoice_number,
            sale.date.strftime('%d/%m/%Y %H:%M'),
            sale.cashier.get_full_name(),
            f"{sale.total_amount:.2f}",
            sale.get_status_display(),
        ])
    # 2) Récupérer le texte et encoder en bytes
    byte_stream = BytesIO(text_stream.getvalue().encode("utf-8"))
    return byte_stream




from io import BytesIO,StringIO 
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
import csv

def generate_employees_pdf(queryset):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Liste des employés")
    y -= 30
    c.setFont("Helvetica", 10)
    cols = ["Nom", "Email", "Rôle", "Date d'inscription"]
    x_positions = [50, 200, 350, 450]
    for x, col in zip(x_positions, cols):
        c.drawString(x, y, col)
    y -= 20
    for emp in queryset:
        if y < 50:
            c.showPage()
            y = height - 50
        values = [
            emp.get_full_name(),
            emp.email,
            emp.get_role_display(),
            emp.date_joined.strftime("%d/%m/%Y"),
        ]
        for x, val in zip(x_positions, values):
            c.drawString(x, y, val)
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer

def generate_employees_excel(queryset):
    wb = Workbook()
    ws = wb.active
    ws.title = "Employés"
    ws.append(["Nom", "Email", "Rôle", "Date d'inscription"])
    for emp in queryset:
        ws.append([
            emp.get_full_name(),
            emp.email,
            emp.get_role_display(),
            emp.date_joined.strftime("%d/%m/%Y"),
        ])
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

def generate_employees_docx(queryset):
    doc = Document()
    doc.add_heading('Liste des employés', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    for idx, title in enumerate(['Nom', 'Email', 'Rôle', "Date d'inscription"]):
        hdr_cells[idx].text = title
    for emp in queryset:
        row_cells = table.add_row().cells
        row_cells[0].text = emp.get_full_name()
        row_cells[1].text = emp.email
        row_cells[2].text = emp.get_role_display()
        row_cells[3].text = emp.date_joined.strftime("%d/%m/%Y")
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

def generate_employees_csv(queryset):
    # 1) Écrire dans un buffer texte
    text_stream = StringIO()
    writer = csv.writer(text_stream, delimiter='\t', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow(["Nom", "Email", "Rôle", "Date d'inscription"])
    for emp in queryset:
        writer.writerow([
            emp.get_full_name(),
            emp.email,
            emp.get_role_display(),
            emp.date_joined.strftime("%d/%m/%Y"),
        ])
    # 2) Récupérer le texte et encoder en bytes
    byte_stream = BytesIO(text_stream.getvalue().encode("utf-8"))
    return byte_stream



from io import BytesIO, StringIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
import csv

def generate_products_pdf(queryset):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Liste des Produits")
    y -= 30
    c.setFont("Helvetica", 10)
    cols = ["ID", "Nom", "Catégorie", "Prix", "Stock"]
    x_positions = [50, 100, 300, 450, 550]
    for x, col in zip(x_positions, cols):
        c.drawString(x, y, col)
    y -= 20
    for prod in queryset:
        if y < 50:
            c.showPage()
            y = height - 50
        values = [
            str(prod.id),
            prod.name,
            prod.category.name if prod.category else "-",
            f"{prod.price:.2f}€",
            str(prod.stock_quantity),
        ]
        for x, val in zip(x_positions, values):
            c.drawString(x, y, val)
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer

def generate_products_excel(queryset):
    wb = Workbook()
    ws = wb.active
    ws.title = "Produits"
    ws.append(['ID', 'Nom', 'Catégorie', 'Prix', 'Stock'])
    for prod in queryset:
        ws.append([
            prod.id,
            prod.name,
            prod.category.name if prod.category else "",
            prod.price,
            prod.stock_quantity,
        ])
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

def generate_products_docx(queryset):
    doc = Document()
    doc.add_heading('Liste des Produits', level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for idx, title in enumerate(['ID', 'Nom', 'Catégorie', 'Prix', 'Stock']):
        hdr[idx].text = title
    for prod in queryset:
        row_cells = table.add_row().cells
        row_cells[0].text = str(prod.id)
        row_cells[1].text = prod.name
        row_cells[2].text = prod.category.name if prod.category else ""
        row_cells[3].text = f"{prod.price:.2f}€"
        row_cells[4].text = str(prod.stock_quantity)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_products_csv(queryset):
    text_stream = StringIO()
    writer = csv.writer(text_stream, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow(['ID', 'Nom', 'Catégorie', 'Prix', 'Stock'])
    for prod in queryset:
        writer.writerow([
            prod.id,
            prod.name,
            prod.category.name if prod.category else "",
            f"{prod.price:.2f}",
            prod.stock_quantity,
        ])
    return BytesIO(text_stream.getvalue().encode('utf-8'))



from io import BytesIO, StringIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
import csv
from docx.shared import Inches

def generate_suppliers_pdf(queryset):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Liste des Fournisseurs")
    y -= 30
    c.setFont("Helvetica", 10)
    cols = ["ID", "Nom Complet", "Email", "Téléphone", "Adresse"]
    # Colonnes élargies pour éviter le retour à la ligne
    x_positions = [50, 150, 300, 450, 600]
    for x, col in zip(x_positions, cols):
        c.drawString(x, y, col)
    y -= 20
    for sup in queryset:
        if y < 50:
            c.showPage()
            y = height - 50
        # Utiliser nom et prénom si disponibles
        full_name = f"{sup.last_name} {sup.first_name}" if hasattr(sup, 'first_name') else sup.name
        values = [
            str(sup.id),
            full_name,
            sup.email or "-",
            sup.phone or "-",
            sup.address or "-",  # remplacez par le champ réel
        ]
        for x, val in zip(x_positions, values):
            c.drawString(x, y, val)
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer

def generate_suppliers_excel(queryset):
    wb = Workbook()
    ws = wb.active
    ws.title = "Fournisseurs"
    ws.append(['ID', 'Nom', 'Email', 'Téléphone', 'Adresse'])
    for sup in queryset:
        ws.append([
            sup.id,
            sup.name,
            sup.email or "",
            sup.phone or "",
            sup.address or "",  # remplacez par le champ réel
        ])
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

def generate_suppliers_docx(queryset):
    doc = Document()
    doc.add_heading('Liste des Fournisseurs', level=1)

    # Crée la table et désactive l'ajustement automatique
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False

    # Définit une largeur fixe pour chaque colonne
    widths = [Inches(0.5), Inches(2.0), Inches(2.5), Inches(1.5), Inches(3.0)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    # En-têtes
    hdr_cells = table.rows[0].cells
    for idx, title in enumerate(['ID', 'Nom Complet', 'Email', 'Téléphone', 'Adresse']):
        hdr_cells[idx].text = title

    # Lignes de données
    for sup in queryset:
        row_cells = table.add_row().cells
        full_name = f"{sup.last_name} {sup.first_name}" if hasattr(sup, 'first_name') else sup.name
        row_cells[0].text = str(sup.id)
        row_cells[1].text = full_name
        row_cells[2].text = sup.email or ""
        row_cells[3].text = sup.phone or ""
        row_cells[4].text = sup.address or ""  # remplacez par le champ réel

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf



def generate_suppliers_csv(queryset):
    text_stream = StringIO()
    writer = csv.writer(text_stream, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow(['ID', 'Nom', 'Email', 'Téléphone', 'Adresse'])
    for sup in queryset:
        writer.writerow([
            sup.id,
            sup.name,
            sup.email or "",
            sup.phone or "",
            sup.address or "",  # remplacez par le champ réel
        ])
    return BytesIO(text_stream.getvalue().encode('utf-8'))



from io import BytesIO, StringIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
import csv

# PDF
def generate_categories_pdf(queryset):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Liste des Catégories")
    y -= 30
    c.setFont("Helvetica", 10)
    cols = ["ID", "Nom", "Description"]
    x_positions = [50, 150, 350]
    for x, col in zip(x_positions, cols):
        c.drawString(x, y, col)
    y -= 20
    for cat in queryset:
        if y < 50:
            c.showPage()
            y = height - 50
        values = [
            str(cat.id),
            cat.name,
            cat.description or "-",
        ]
        for x, val in zip(x_positions, values):
            c.drawString(x, y, val)
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer

# Excel
def generate_categories_excel(queryset):
    wb = Workbook()
    ws = wb.active
    ws.title = "Catégories"
    ws.append(['ID', 'Nom', 'Description'])
    for cat in queryset:
        ws.append([
            cat.id,
            cat.name,
            cat.description or "",
        ])
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

# Word
from io import BytesIO
from docx import Document
from docx.shared import Inches

def generate_categories_docx(queryset):
    doc = Document()
    doc.add_heading('Liste des Catégories', level=1)

    # Crée la table et désactive l'ajustement automatique
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False

    # Définir largueurs de colonnes
    widths = [Inches(0.5), Inches(2.0), Inches(3.5)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    # En-têtes
    hdr_cells = table.rows[0].cells
    for idx, title in enumerate(['ID', 'Nom', 'Description']):
        hdr_cells[idx].text = title

    # Lignes de données
    for cat in queryset:
        row_cells = table.add_row().cells
        row_cells[0].text = str(cat.id)
        row_cells[1].text = cat.name or ""
        row_cells[2].text = cat.description or ""

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# CSV/TSV
def generate_categories_csv(queryset):
    text_stream = StringIO()
    writer = csv.writer(text_stream, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow(['ID', 'Nom', 'Description'])
    for cat in queryset:
        writer.writerow([
            cat.id,
            cat.name,
            cat.description or "",
        ])
    return BytesIO(text_stream.getvalue().encode('utf-8'))



from io import BytesIO, StringIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches
import csv
from django.utils import timezone

# 1. Ventes (chiffre d'affaires)
def generate_sales_report_pdf(data):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    title = f"Rapport des Ventes – {timezone.now().date()}"
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, title)
    y -= 30
    c.setFont("Helvetica", 10)
    cols = ["Date", "Total TTC", "Nombre de transactions"]
    x_positions = [50, 200, 400]
    for x, col in zip(x_positions, cols):
        c.drawString(x, y, col)
    y -= 20
    for row in data:
        if y < 50:
            c.showPage()
            y = height - 50
        values = [
            row["date"].strftime("%Y-%m-%d"),
            f"{row['total']:.2f}€",
            str(row["count"])
        ]
        for x, val in zip(x_positions, values):
            c.drawString(x, y, val)
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer

def generate_sales_report_excel(data):
    wb = Workbook()
    ws = wb.active
    ws.title = "Ventes"
    ws.append(["Date", "Total TTC", "Nb transactions"])
    for row in data:
        date_str = row["date"].strftime("%Y-%m-%d")
        ws.append([date_str, row["total"], row["count"]])
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

def generate_sales_report_docx(data):
    doc = Document()
    doc.add_heading("Rapport des Ventes", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(1.5), Inches(2.0), Inches(1.5)]
    for idx, w in enumerate(widths):
        table.columns[idx].width = w
    hdr = table.rows[0].cells
    for idx, title in enumerate(["Date", "Total TTC", "Nb transactions"]):
        hdr[idx].text = title
    for row in data:
        cells = table.add_row().cells
        cells[0].text = row["date"].strftime("%Y-%m-%d")
        cells[1].text = f"{row['total']:.2f}€"
        cells[2].text = str(row["count"])
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_sales_report_csv(data):
    text_stream = StringIO()
    writer = csv.writer(text_stream, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow(["Date", "Total TTC", "Nb transactions"])
    for row in data:
        date_str = row["date"].strftime("%Y-%m-%d")
        writer.writerow([date_str, f"{row['total']:.2f}", row["count"]])
    return BytesIO(text_stream.getvalue().encode("utf-8"))

# 2. Stocks (quantités restantes)
def generate_stock_report_pdf(data):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 50, "Rapport des Stocks")
    y = height - 80
    c.setFont("Helvetica", 10)
    cols = ["Produit", "Catégorie", "En stock"]
    x_positions = [50, 250, 450]
    for x, col in zip(x_positions, cols):
        c.drawString(x, y, col)
    y -= 20
    for row in data:
        if y < 50:
            c.showPage()
            y = height - 50
        values = [row["name"], row["category"], str(row["stock"])]
        for x, val in zip(x_positions, values):
            c.drawString(x, y, val)
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer

def generate_stock_report_excel(data):
    wb = Workbook()
    ws = wb.active
    ws.title = "Stocks"
    ws.append(["Produit", "Catégorie", "En stock"])
    for row in data:
        ws.append([row["name"], row["category"], row["stock"]])
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

def generate_stock_report_docx(data):
    doc = Document()
    doc.add_heading("Rapport des Stocks", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(2.5), Inches(2.0), Inches(1.0)]
    for idx, w in enumerate(widths):
        table.columns[idx].width = w
    hdr = table.rows[0].cells
    for idx, title in enumerate(["Produit", "Catégorie", "En stock"]):
        hdr[idx].text = title
    for row in data:
        cells = table.add_row().cells
        cells[0].text = row["name"]
        cells[1].text = row["category"]
        cells[2].text = str(row["stock"])
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_stock_report_csv(data):
    text_stream = StringIO()
    writer = csv.writer(text_stream, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow(["Produit", "Catégorie", "En stock"])
    for row in data:
        writer.writerow([row["name"], row["category"], row["stock"]])
    return BytesIO(text_stream.getvalue().encode("utf-8"))
